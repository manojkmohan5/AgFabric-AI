"""Text extraction from uploaded documents and images.

Pure functions over bytes, so they are testable without storage or a database.
Every parser is pure Python — no tesseract, no poppler, no native toolchain.

Images go through `ocr.py` (vision API). Scanned PDFs are handled by pulling the
embedded images out with pypdf and OCRing those, which avoids rasterising and so
avoids a PDF-rendering dependency entirely.

`extract` returns (text, note). The note explains anything the caller should
surface — OCR was used, some pages held no text, only the first N pages were read.
"""

import csv
import io
import logging
from pathlib import PurePosixPath

# Only these are accepted. An allowlist, so an unrecognised type is rejected
# rather than guessed at.
from . import ocr
from .config import settings

logger = logging.getLogger(__name__)

DOCUMENT_SUFFIXES = frozenset({".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"})
ALLOWED_SUFFIXES = DOCUMENT_SUFFIXES | ocr.IMAGE_SUFFIXES
MAX_SHEET_ROWS = 5_000
# Pages to OCR from a scanned PDF. Each is a billed vision call, so bounded.
MAX_PDF_OCR_PAGES = 5


class ExtractionError(Exception):
    """The file matched an allowed suffix but could not be parsed."""


def safe_suffix(filename: str) -> str:
    """Lowercased extension of `filename`, with any directory path discarded.

    Handles both separators, so neither `../../etc/passwd` nor
    `..\\windows\\system32` can smuggle a path through.
    """
    basename = PurePosixPath(filename.replace("\\", "/")).name
    return PurePosixPath(basename).suffix.lower()


def extract(data: bytes, filename: str) -> tuple[str, str | None]:
    """Return (text, note). `note` is None unless something needs explaining."""
    suffix = safe_suffix(filename)
    if suffix not in ALLOWED_SUFFIXES:
        raise ExtractionError(f"unsupported file type {suffix or '(none)'}")
    try:
        if suffix in ocr.IMAGE_SUFFIXES:
            return _image(data, suffix)
        if suffix == ".pdf":
            return _pdf(data)
        if suffix == ".docx":
            return _docx(data), None
        if suffix == ".xlsx":
            return _xlsx(data), None
        if suffix == ".csv":
            return _csv(data), None
        return _plain(data), None
    except ExtractionError:
        raise
    except Exception as exc:  # a corrupt upload must not 500
        raise ExtractionError(f"could not parse {suffix} file: {exc}") from exc


def _image(data: bytes, suffix: str) -> tuple[str, str | None]:
    """OCR an uploaded image."""
    if len(data) > settings.max_image_bytes:
        raise ExtractionError(
            f"image exceeds the {settings.max_image_bytes // (1024 * 1024)}MB limit "
            "for OCR; downscale it first"
        )
    try:
        mime = ocr.sniff_image(data, suffix)
        text = ocr.get_reader().read(data, mime)
    except ocr.OCRError as exc:
        raise ExtractionError(str(exc)) from exc

    reader = ocr.get_reader()
    if not text:
        return "", f"No legible text found by {reader.name} OCR."
    return text, f"Text recovered by {reader.name} OCR ({reader.model})."


def _pdf(data: bytes) -> tuple[str, str | None]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise ExtractionError("PDF is encrypted")
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(p for p in pages if p)

    if not text:
        # A scanned PDF holds its content as embedded images. pypdf can pull those
        # out without rasterising, so OCR works here with no new dependency —
        # no poppler, no pdfium.
        recovered, note = _pdf_via_ocr(reader)
        if recovered:
            return recovered, note
        return "", (
            f"No embedded text in {len(pages)} page(s) and no images could be "
            f"read. {note or ''}".strip()
        )

    empty = sum(1 for p in pages if not p)
    note = f"{empty} of {len(pages)} pages held no embedded text." if empty else None
    return text, note


def _pdf_via_ocr(reader) -> tuple[str, str | None]:  # noqa: ANN001
    """OCR the images embedded in a scanned PDF, up to MAX_PDF_OCR_PAGES."""
    engine = ocr.get_reader()
    parts: list[str] = []
    attempted = 0
    for index, page in enumerate(reader.pages[:MAX_PDF_OCR_PAGES], start=1):
        try:
            images = list(page.images)
        except Exception as exc:
            # A malformed image stream must not sink the whole upload, but it
            # should not vanish either.
            logger.warning("could not read images on PDF page %s: %s", index, exc)
            continue
        for image in images:
            suffix = ocr.IMAGE_SUFFIXES & {f".{image.name.rsplit('.', 1)[-1].lower()}"}
            if not suffix:
                continue
            attempted += 1
            try:
                mime = ocr.sniff_image(image.data, next(iter(suffix)))
                page_text = engine.read(image.data, mime)
            except ocr.OCRError:
                continue
            if page_text:
                parts.append(f"[page {index}]\n{page_text}")
            break  # one image per page is enough for a scan

    if not parts:
        return "", f"tried OCR on {attempted} embedded image(s), none yielded text"
    truncated = len(reader.pages) > MAX_PDF_OCR_PAGES
    note = f"Scanned PDF read by {engine.name} OCR from embedded images."
    if truncated:
        note += f" Only the first {MAX_PDF_OCR_PAGES} of {len(reader.pages)} pages."
    return "\n\n".join(parts), note


def _docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    # read_only + values_only keeps memory flat on wide sheets.
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet in workbook.worksheets:
            rows = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= MAX_SHEET_ROWS:
                    rows.append(f"... truncated at {MAX_SHEET_ROWS} rows")
                    break
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"# {sheet.title}\n" + "\n".join(rows))
        return "\n\n".join(parts)
    finally:
        workbook.close()


def _csv(data: bytes) -> str:
    text = _plain(data)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel  # single column, or too small to sniff
    rows = [
        " | ".join(cell.strip() for cell in row)
        for row in csv.reader(io.StringIO(text), dialect)
        if any(cell.strip() for cell in row)
    ]
    return "\n".join(rows)


def _plain(data: bytes) -> str:
    # Uploads are untrusted bytes; replace undecodable sequences rather than
    # raising, so one bad byte does not reject an otherwise fine document.
    return data.decode("utf-8", errors="replace")
